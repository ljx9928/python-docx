from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def save_link():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A hyperlink paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    p._insert_hyperlink_after('www.baidu.com','hyperlink')

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    document.save('hltest.docx')


def iter_hyperlink_rels(rels):
    for rel in rels.values():
        if rel.reltype == RT.HYPERLINK:
            yield rel

def is_hyperlink_p(p):
    if not p._element is None:
        if not p._element.hl_lst.__len__() is 0:
            if p._element.hl_lst[0].rId in hls.keys():
                return True
    return False


def get_link():
    document = Document("hltest.docx")

    hls = {}
    rels = document.part.rels
    for rel in iter_hyperlink_rels(rels):
        hls[rel.rId] = rel._target

    for p in document.paragraphs:
        if not p._element is None:
            print("P:",p.text.strip())
            if not p._element.hl_lst.__len__() is 0:
                if p._element.hl_lst[0].rId in hls.keys():
                    print( p.text, ":", hls[p._element.hl_lst[0].rId])

if __name__ == '__main__':
    # save_link()
    get_link()